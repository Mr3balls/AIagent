from __future__ import annotations

from datetime import datetime, timezone
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

DOCX_MEDIA_TYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
REPORTS_DIR = Path("/app/storage/reports")


def _set_default_style(document: Document) -> None:
    normal_style = document.styles["Normal"]
    normal_style.font.name = "Times New Roman"
    normal_style.font.size = Pt(12)


def _add_title(document: Document, title: str) -> None:
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(14)


def _add_section_heading(document: Document, text: str) -> None:
    p = document.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)


def _add_paragraph(document: Document, text: str) -> None:
    if not text:
        return
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)


def _add_bullets(document: Document, items: list[str]) -> None:
    for item in items:
        if not item:
            continue
        p = document.add_paragraph(style="List Bullet")
        run = p.add_run(item)
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)


def _format_general_info(context: dict[str, Any]) -> list[str]:
    general = context.get("general_information") or {}

    result: list[str] = []

    if context.get("title"):
        result.append(f'Наименование проекта: {context["title"]}.')
    if context.get("customer_name"):
        result.append(f'Заказчик: {context["customer_name"]}.')
    if general.get("project_type"):
        result.append(f'Тип проекта: {general["project_type"]}.')
    if general.get("implementation_location"):
        result.append(f'Место реализации: {general["implementation_location"]}.')
    if general.get("implementation_days"):
        result.append(f'Срок реализации: {general["implementation_days"]} календарных дней.')
    if general.get("budget"):
        result.append(f'Бюджет: {general["budget"]}.')

    return result


def generate_tender_report_docx(
    *,
    tender_id: str,
    context: dict[str, Any],
) -> dict[str, Any]:
    REPORTS_DIR.mkdir(parents=True, exist_ok=True)

    filename = f"tender_analysis_{tender_id}.docx"
    path = REPORTS_DIR / filename

    document = Document()
    _set_default_style(document)

    _add_title(document, "Аналитический отчет по тендерной документации")
    document.add_paragraph()

    _add_section_heading(document, "1. Executive Summary")
    _add_paragraph(document, context.get("executive_summary", ""))

    _add_section_heading(document, "2. Общие сведения")
    _add_bullets(document, _format_general_info(context))

    _add_section_heading(document, "3. Краткое содержание технического задания")
    _add_paragraph(document, context.get("technical_specification_summary", ""))

    _add_section_heading(document, "4. Ключевые требования")
    _add_bullets(document, context.get("key_requirements", []))

    _add_section_heading(document, "5. Полнота документации")
    _add_paragraph(document, context.get("documentation_completeness_text", ""))

    _add_section_heading(document, "6. Признаки возможной заточки")
    _add_paragraph(document, context.get("tailoring_text", ""))

    _add_section_heading(document, "7. Анализ сроков")
    _add_paragraph(document, context.get("timeline_analysis", ""))

    _add_section_heading(document, "8. Анализ рисков")
    _add_paragraph(document, context.get("risks_text", ""))

    _add_section_heading(document, "9. Уточняющие вопросы заказчику")
    _add_paragraph(document, context.get("questions_text", ""))

    _add_section_heading(document, "10. Предварительное решение")
    decision = context.get("decision", "не определено")
    risk_score = context.get("risk_score", "не определен")
    _add_paragraph(
        document,
        f'По итогам автоматизированного анализа итоговая категория решения определена как "{decision}", '
        f"итоговый риск-скор составляет {risk_score}.",
    )

    _add_section_heading(document, "11. Финальное заключение")
    _add_paragraph(document, context.get("conclusion", ""))

    document.save(path)

    return {
        "path": str(path),
        "filename": filename,
        "generated_at": datetime.now(timezone.utc),
        "media_type": DOCX_MEDIA_TYPE,
    }