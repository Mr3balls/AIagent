from __future__ import annotations

from pathlib import Path
from typing import Any

from docx import Document

from app.services.parsers.document_parser import DocumentParser


class DocxParser:
    def parse(self, file_path: str | Path) -> dict[str, Any]:
        path = Path(file_path)
        document = Document(path)

        paragraphs_text: list[str] = []
        sections: list[dict[str, Any]] = []
        tables: list[dict[str, Any]] = []

        current_section_title: str | None = None
        current_section_content: list[str] = []

        for paragraph in document.paragraphs:
            text = paragraph.text.strip()
            if not text:
                continue

            style_name = paragraph.style.name.lower() if paragraph.style and paragraph.style.name else ""

            paragraphs_text.append(text)

            if "heading" in style_name:
                if current_section_content:
                    sections.append(
                        {
                            "index": len(sections) + 1,
                            "title": current_section_title,
                            "content": "\n".join(current_section_content).strip(),
                        }
                    )
                    current_section_content = []

                current_section_title = text
            else:
                current_section_content.append(text)

        if current_section_content:
            sections.append(
                {
                    "index": len(sections) + 1,
                    "title": current_section_title,
                    "content": "\n".join(current_section_content).strip(),
                }
            )

        for table_index, table in enumerate(document.tables, start=1):
            rows: list[list[str]] = []
            for row in table.rows:
                row_values = [cell.text.strip() for cell in row.cells]
                if any(row_values):
                    rows.append(row_values)

            if rows:
                tables.append(
                    {
                        "index": table_index,
                        "title": f"DOCX table {table_index}",
                        "rows": rows,
                    }
                )

        merged_text = "\n\n".join(paragraphs_text).strip()
        normalized_text = DocumentParser.normalize_text(merged_text)

        if not sections:
            sections = DocumentParser.split_text_into_blocks(normalized_text)

        return {
            "text": normalized_text,
            "tables": tables,
            "sections": sections,
        }